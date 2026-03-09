"""
Resume Parser - Extract text from PDF, DOCX, and TXT files
"""
import os
from pathlib import Path
from typing import Optional
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def parse_resume(file_path: str, file_content: bytes = None) -> Optional[str]:
    """
    Extract text from resume files (PDF, DOCX, TXT)
    
    Args:
        file_path: Path to the resume file or filename with extension
        file_content: Optional bytes content of the file
        
    Returns:
        Extracted text as string, or None if parsing fails
    """
    file_extension = Path(file_path).suffix.lower()
    
    try:
        if file_extension == '.txt':
            return _parse_txt(file_content)
        elif file_extension == '.pdf':
            return _parse_pdf(file_content)
        elif file_extension in ['.docx', '.doc']:
            return _parse_docx(file_content)
        else:
            logger.error(f"Unsupported file format: {file_extension}")
            return None
    except Exception as e:
        logger.error(f"Error parsing resume: {e}")
        return None


def _parse_txt(file_content: bytes) -> str:
    """Parse plain text file"""
    try:
        # Try UTF-8 first, fallback to latin-1
        text = file_content.decode('utf-8')
    except UnicodeDecodeError:
        text = file_content.decode('latin-1', errors='ignore')
    
    return text.strip()


def _parse_pdf(file_content: bytes) -> str:
    """Parse PDF file using PyPDF2"""
    try:
        from PyPDF2 import PdfReader
        from io import BytesIO
        
        pdf_file = BytesIO(file_content)
        reader = PdfReader(pdf_file)
        
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        
        return text.strip()
    except ImportError:
        logger.error("PyPDF2 not installed. Install with: pip install PyPDF2")
        return None
    except Exception as e:
        logger.error(f"Error parsing PDF: {e}")
        return None


def _parse_docx(file_content: bytes) -> str:
    """Parse DOCX file using python-docx"""
    try:
        from docx import Document
        from io import BytesIO
        
        docx_file = BytesIO(file_content)
        doc = Document(docx_file)
        
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
            text += "\n"
        
        return text.strip()
    except ImportError:
        logger.error("python-docx not installed. Install with: pip install python-docx")
        return None
    except Exception as e:
        logger.error(f"Error parsing DOCX: {e}")
        return None


def clean_text(text: str) -> str:
    """
    Clean extracted text by removing extra whitespace and special characters
    
    Args:
        text: Raw extracted text
        
    Returns:
        Cleaned text
    """
    if not text:
        return ""
    
    # Replace multiple spaces with single space
    import re
    text = re.sub(r'\s+', ' ', text)
    
    # Remove non-printable characters
    text = ''.join(char for char in text if char.isprintable() or char in ['\n', '\t'])
    
    return text.strip()


if __name__ == "__main__":
    # Test the parser
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python resume_parser.py <file_path>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        sys.exit(1)
    
    with open(file_path, 'rb') as f:
        content = f.read()
    
    text = parse_resume(file_path, content)
    
    if text:
        print(f"Successfully extracted {len(text)} characters")
        print("\n--- First 500 characters ---")
        print(text[:500])
    else:
        print("Failed to extract text")
